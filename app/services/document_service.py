from __future__ import annotations

import html
import io
import json
import logging
import os
import re
import uuid
from pathlib import Path
from typing import Any

import httpx
import pdfplumber
from bs4 import BeautifulSoup
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph
from pypdf import PdfReader
from reportlab.lib.colors import black
from reportlab.lib.enums import TA_CENTER, TA_RIGHT
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable,
    ListFlowable,
    ListItem,
    Paragraph as PdfParagraph,
    SimpleDocTemplate,
    Table,
    TableStyle,
)

from app.models import Achievement, Course, ParsedJob, Profile, Project, Role, Skill, Template
from app.schemas import ParsedJobSchema
from app.services.claude_service import call_claude
from app.services.error_handlers import ClaudeAPIError, PDFExtractionError, TemplateError, URLFetchError
from app.services.prompt_templates import (
    build_cover_letter_customization_prompt,
    build_cold_email_prompt,
    build_cv_customization_prompt,
    build_cv_structured_prompt,
)

logger = logging.getLogger(__name__)


def _load_env_file() -> None:
    env_path = Path(__file__).resolve().parents[2] / ".env"
    if env_path.exists():
        for line in env_path.read_text().splitlines():
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, value = line.split("=", 1)
            os.environ.setdefault(key.strip(), value.strip().strip('"').strip("'"))


_load_env_file()


def extract_pdf_text(contents: bytes) -> str:
    """Extract text from a PDF's raw bytes. Raises PDFExtractionError, never crashes."""
    try:
        with pdfplumber.open(io.BytesIO(contents)) as pdf:
            if not pdf.pages:
                raise PDFExtractionError("PDF contains no readable text")

            page_texts: list[str] = []
            has_text = False
            has_images = False
            for page in pdf.pages:
                page_text = page.extract_text() or ""
                if page_text.strip():
                    has_text = True
                if page.images:
                    has_images = True
                page_texts.append(page_text)

            if not has_text:
                if has_images:
                    raise PDFExtractionError("PDF is image-only; OCR not supported")
                raise PDFExtractionError("PDF contains no readable text")

            return "\n".join(page_texts).strip()
    except PDFExtractionError:
        raise
    except Exception as exc:
        logger.error("Failed to open PDF: %s", exc)
        raise PDFExtractionError("Invalid PDF file") from exc


def extract_url_text(url: str) -> str:
    """Fetch a URL and extract its visible text. Raises URLFetchError, never crashes."""
    try:
        response = httpx.get(url, timeout=10.0, follow_redirects=True, headers={"User-Agent": "Mozilla/5.0"})
    except httpx.TimeoutException as exc:
        logger.error("Timed out fetching URL %s: %s", url, exc)
        raise URLFetchError("Request timed out") from exc
    except httpx.ConnectError as exc:
        logger.error("Could not connect to URL %s: %s", url, exc)
        raise URLFetchError("Could not reach URL") from exc
    except httpx.RequestError as exc:
        logger.error("Request error fetching URL %s: %s", url, exc)
        raise URLFetchError("Could not reach URL") from exc

    if response.status_code != 200:
        logger.error("URL %s returned HTTP %s", url, response.status_code)
        raise URLFetchError(f"HTTP {response.status_code}")

    soup = BeautifulSoup(response.text, "html.parser")
    for tag in soup(["script", "style", "noscript"]):
        tag.decompose()
    return " ".join(soup.get_text(separator=" ").split())


def _open_template_or_raise(template_path: str) -> Document:
    if not template_path or not Path(template_path).exists():
        raise TemplateError(f"Template not found: {template_path!r}")
    try:
        return Document(template_path)
    except Exception as exc:
        raise TemplateError(f"Template is corrupted: {exc}") from exc


def load_template_document(template_path: str) -> Document:
    """Load a template DOCX, falling back to a blank Document on any failure.

    Always returns a valid Document — never raises. Missing or corrupted templates
    are logged as warnings and silently replaced with a blank starting point.
    """
    try:
        return _open_template_or_raise(template_path)
    except TemplateError as exc:
        logger.warning("%s — using blank document fallback", exc)
        return Document()


def _extract_text_from_pdf(path: str) -> str:
    try:
        reader = PdfReader(path)
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    except Exception as exc:
        logger.warning("Failed to extract text from PDF %r: %s", path, exc)
        return ""


def load_template_text(template_path: str) -> str:
    """Return plain text from a template file regardless of source format
    (PDF or DOCX). Never raises; returns "" for empty/missing paths."""
    if not template_path or not Path(template_path).exists():
        return ""
    if template_path.lower().endswith(".pdf"):
        return _extract_text_from_pdf(template_path)
    template_doc = load_template_document(template_path)
    return "\n\n".join(
        paragraph.text.strip() for paragraph in template_doc.paragraphs if paragraph.text.strip()
    )


def _build_user_profile_json(
    profile: Profile | None,
    roles: list[Role],
    projects: list[Project],
    skills: list[Skill],
    courses: list[Course] | None = None,
    achievements: list[Achievement] | None = None,
) -> dict[str, Any]:
    return {
        "profile": {
            "name": profile.name if profile else "",
            "email": profile.email if profile else "",
            "phone": profile.phone if profile else "",
            "linkedin": profile.linkedin if profile else "",
            "portfolio": profile.portfolio_url if profile else "",
            "github": profile.github_url if profile else "",
            "location": profile.location if profile else "",
        },
        "roles": [
            {
                "title": role.title,
                "company": role.company,
                "startDate": role.start_date,
                "endDate": role.end_date,
                "description": role.description,
                "achievements": role.achievements or [],
                "metrics": role.metrics,
            }
            for role in roles
        ],
        "projects": [
            {
                "title": project.title,
                "description": project.description,
                "technologies": project.technologies or [],
                "outcomes": project.outcomes or [],
                "metrics": project.metrics,
                "reflection": project.reflection,
            }
            for project in projects
        ],
        "courses": [
            {
                "name": course.name,
                "provider": course.provider,
                "dateCompleted": course.date_completed,
                "learnings": course.learnings,
            }
            for course in (courses or [])
        ],
        "achievements": [
            {
                "title": achievement.title,
                "type": achievement.type,
                "date": achievement.date,
                "details": achievement.details,
            }
            for achievement in (achievements or [])
        ],
        "skills": [
            {
                "name": skill.name,
                "category": skill.category,
                "proficiency": skill.proficiency,
                "relatedProjectIds": skill.related_project_ids or [],
                "relatedRoleIds": skill.related_role_ids or [],
            }
            for skill in skills
        ],
    }


def _clean_json_payload(text: str) -> str:
    cleaned = text.strip()
    if cleaned.startswith("```json"):
        cleaned = cleaned[len("```json"):]
    if cleaned.startswith("```"):
        cleaned = cleaned[3:]
    if cleaned.endswith("```"):
        cleaned = cleaned[:-3]
    return cleaned.strip()


def _clean_text_payload(text: str) -> str:
    cleaned = _clean_json_payload(text)
    cleaned = cleaned.replace("Subject:", "").strip()
    cleaned = re.sub(r"\n{3,}", "\n\n", cleaned)
    return cleaned.strip()


def _looks_like_structure(text: str) -> bool:
    lowered = text.lower().strip()
    if not lowered:
        return False
    if any(marker in lowered for marker in ["@", "www.", ".com", "http", "phone", "email", "address"]):
        return True
    if any(marker in lowered for marker in ["sincerely", "best", "regards", "kind regards", "thanks", "thank you", "signature"]):
        return True
    return len(text.split()) <= 4 and not any(marker in lowered for marker in [".", "!", "?"])


def _copy_paragraph(src_paragraph: Paragraph, dst_doc: Document) -> None:
    new_paragraph = dst_doc.add_paragraph()
    new_paragraph.style = src_paragraph.style
    for run in src_paragraph.runs:
        new_run = new_paragraph.add_run(run.text)
        new_run.bold = run.bold
        new_run.italic = run.italic
        new_run.underline = run.underline
        new_run.font.name = run.font.name
        new_run.font.size = run.font.size


def _build_fallback_cover_letter(user_profile_json: dict, parsed_job: dict) -> str:
    """Deterministic, metrics-aware fallback used when Claude is unavailable.

    Ranks roles/projects by overlap with the job's requiredSkills/keywords, same as
    the CV fallback, so the 2-3 achievements cited are the most relevant ones on
    file rather than whichever happened to be entered first.
    """
    company = parsed_job.get("company") or "the company"
    role_title = parsed_job.get("title") or "this role"
    job_terms = {
        term.lower()
        for term in (parsed_job.get("requiredSkills", []) + parsed_job.get("keywords", []))
    }

    ranked_roles = sorted(
        user_profile_json.get("roles", []),
        key=lambda role: _relevance_score(
            [role.get("title", ""), role.get("company", ""), role.get("description", "")], job_terms
        ),
        reverse=True,
    )
    ranked_projects = sorted(
        user_profile_json.get("projects", []),
        key=lambda project: _relevance_score(
            [project.get("title", ""), project.get("description", "")]
            + (project.get("technologies") or [])
            + (project.get("outcomes") or []),
            job_terms,
        ),
        reverse=True,
    )

    paragraphs = [
        f"At {company}, your work on {role_title} stood out because I have built comparable systems for high-throughput, reliability-focused products."
    ]

    if ranked_roles:
        top_role = ranked_roles[0]
        metric = top_role.get("metrics") or next(iter(top_role.get("achievements") or []), None) or "drove measurable impact"
        paragraphs.append(
            f"As {top_role.get('title', 'an engineer')} at {top_role.get('company', 'my last role')}, I achieved {metric}."
        )

    if ranked_projects:
        top_project = ranked_projects[0]
        metric = top_project.get("metrics") or next(iter(top_project.get("outcomes") or []), None) or "shipped a production system"
        paragraphs.append(f"I also built {top_project.get('title', 'a project')}, resulting in {metric}.")

    achievements = user_profile_json.get("achievements") or []
    if achievements:
        achievement = achievements[0]
        detail = achievement.get("details") or "drove adoption"
        paragraphs.append(f"I was also recognized for {achievement.get('title', 'my work')}: {detail}")

    paragraphs.append(
        f"I would love to bring that experience to {company} and help strengthen the systems behind {role_title}. "
        "Would you have 15 minutes Thursday or Friday to discuss how I could support the team?"
    )
    return "\n\n".join(paragraphs)


def _build_fallback_cold_email(user_profile_json: dict, parsed_job: dict) -> str:
    company = parsed_job.get("company") or "your team"
    role = parsed_job.get("title") or "this role"
    metric = ""
    for project_entry in user_profile_json.get("projects", []):
        if project_entry.get("metrics"):
            metric = project_entry.get("metrics")
            break
    if not metric:
        for role_entry in user_profile_json.get("roles", []):
            if role_entry.get("metrics"):
                metric = role_entry.get("metrics")
                break
    if not metric:
        metric = "28% better accuracy"

    return (
        f"Hi there, I was reading about {company}'s work around {role} and it immediately stood out because I recently built a comparable platform that delivered {metric}. "
        f"I enjoy solving the kind of reliability and scale problems that show up in high-throughput products, and I would be excited to learn more about the team's roadmap. "
        f"Would you have 15 minutes Thursday or Friday to discuss the work?"
    )


def _parse_bullets(text: str) -> list[str]:
    cleaned = _clean_json_payload(text)
    parts = cleaned.split("•") if "•" in cleaned else cleaned.splitlines()
    bullets = [part.strip(" \t\n-*") for part in parts]
    return [bullet for bullet in bullets if bullet]


def _relevance_score(text_fields: list[str], job_terms: set[str]) -> int:
    haystack = " ".join(text_fields).lower()
    return sum(1 for term in job_terms if term in haystack)


def _fallback_cv_bullets(user_profile_json: dict, parsed_job: dict) -> list[str]:
    """Deterministic, metrics-aware fallback used when Claude is unavailable.

    Ranks roles/projects by overlap with the job's requiredSkills/keywords so the
    most relevant experience is chosen first, same intent as step 1 of the prompt.
    """
    job_terms = {
        term.lower()
        for term in (parsed_job.get("requiredSkills", []) + parsed_job.get("keywords", []))
    }
    bullets: list[str] = []

    ranked_roles = sorted(
        user_profile_json.get("roles", []),
        key=lambda role: _relevance_score(
            [role.get("title", ""), role.get("company", ""), role.get("description", "")], job_terms
        ),
        reverse=True,
    )
    ranked_projects = sorted(
        user_profile_json.get("projects", []),
        key=lambda project: _relevance_score(
            [project.get("title", ""), project.get("description", "")]
            + (project.get("technologies") or [])
            + (project.get("outcomes") or []),
            job_terms,
        ),
        reverse=True,
    )

    for role in ranked_roles[:2]:
        metric = role.get("metrics") or next(iter(role.get("achievements") or []), None) or "drove adoption"
        matched = sorted(t for t in job_terms if t in (role.get("description") or "").lower())
        keyword_note = f" using {', '.join(matched[:2])}" if matched else ""
        bullets.append(
            f"Led {role.get('title', 'role')} work at {role.get('company', 'the company')}{keyword_note}, achieving {metric}."
        )

    for project in ranked_projects[:2]:
        metric = project.get("metrics") or next(iter(project.get("outcomes") or []), None) or "designed system for production use"
        techs = project.get("technologies") or []
        tech_note = f" with {', '.join(techs[:2])}" if techs else ""
        bullets.append(f"Built {project.get('title', 'project')}{tech_note}, resulting in {metric}.")

    for achievement in user_profile_json.get("achievements", [])[:1]:
        detail = achievement.get("details") or "drove adoption"
        bullets.append(f"Recognized for {achievement.get('title', 'achievement')}: {detail}.")

    return bullets or ["Drove adoption of new engineering practices across the team."]


def customize_cv_text(user_profile_json: dict, parsed_job: dict) -> list[str]:
    """Never raises: any Claude failure falls back to deterministic, metrics-aware bullets."""
    prompt = build_cv_customization_prompt(user_profile_json, parsed_job)
    try:
        raw_text = call_claude(prompt, max_tokens=1200)
        bullets = _parse_bullets(raw_text)
        return bullets or _fallback_cv_bullets(user_profile_json, parsed_job)
    except ClaudeAPIError as exc:
        logger.warning("customize_cv_text falling back to deterministic bullets: %s", exc)
        return _fallback_cv_bullets(user_profile_json, parsed_job)


def _fallback_cv_structured(user_profile_json: dict) -> dict:
    """Deterministic structured CV built straight from profile data, no Claude tailoring."""
    profile = user_profile_json.get("profile", {}) or {}
    contact = " | ".join(
        part for part in (
            profile.get("location"),
            profile.get("phone"),
            profile.get("email"),
            profile.get("linkedin"),
            profile.get("portfolio"),
            profile.get("github"),
        ) if part
    )

    skills_by_category: dict[str, list[str]] = {}
    for skill in user_profile_json.get("skills", []):
        category = skill.get("category") or "Skills"
        skills_by_category.setdefault(category, []).append(skill.get("name", ""))

    experience = [
        {
            "role": role.get("title", ""),
            "company": role.get("company", ""),
            "dates": f"{role.get('startDate', '')} - {role.get('endDate', '')}",
            "location": "",
            "bullets": role.get("achievements") or ([role["description"]] if role.get("description") else []),
        }
        for role in user_profile_json.get("roles", [])
    ]

    projects = [
        {"name": project.get("title", ""), "description": project.get("description", "")}
        for project in user_profile_json.get("projects", [])
    ]

    return {
        "name": profile.get("name", ""),
        "contact": contact,
        "education": [],
        "skills": skills_by_category,
        "experience": experience,
        "projects": projects,
        "leadership": [],
    }


_CV_STRUCTURED_KEYS = {"name", "contact", "education", "skills", "experience", "projects", "leadership"}


def generate_cv_structured(user_profile_json: dict, parsed_job: dict, template_text: str) -> dict:
    """Never raises: any Claude failure or malformed/incomplete JSON falls back to a
    deterministic structure built directly from user_profile_json."""
    prompt = build_cv_structured_prompt(user_profile_json, parsed_job, template_text)
    try:
        raw_text = call_claude(prompt, max_tokens=20000)
        content = _clean_json_payload(raw_text)
        data = json.loads(content)
        if not isinstance(data, dict) or not _CV_STRUCTURED_KEYS.issubset(data.keys()):
            raise ValueError("Claude response missing required keys")
        return data
    except (ClaudeAPIError, json.JSONDecodeError, ValueError, TypeError) as exc:
        logger.warning("generate_cv_structured falling back to deterministic structure: %s", exc)
        return _fallback_cv_structured(user_profile_json)


def render_cv_html(cv: dict) -> str:
    """Render a structured CV dict (the shape produced by generate_cv_structured)
    as a single self-contained HTML string — inline <style>, no external
    stylesheet or JS, safe to drop into dangerouslySetInnerHTML or an iframe
    srcDoc. Every string value is html.escape'd before insertion."""

    def esc(value: Any) -> str:
        return html.escape(str(value or ""))

    name = esc(cv.get("name"))
    contact = esc(cv.get("contact"))

    sections_html: list[str] = []

    education = cv.get("education") or []
    if education:
        rows = []
        for entry in education:
            school = esc(entry.get("school"))
            degree = esc(entry.get("degree"))
            dates = esc(entry.get("dates"))
            location = esc(entry.get("location"))
            title = f"{degree}, {school}" if degree and school else (degree or school)
            meta = " · ".join(part for part in (dates, location) if part)
            rows.append(f"""
            <div class="entry">
              <div class="entry-header">
                <span class="entry-title">{title}</span>
                <span class="entry-meta">{meta}</span>
              </div>
            </div>
            """)
        sections_html.append(f"""
        <section>
          <h2>Education</h2>
          {''.join(rows)}
        </section>
        """)

    experience = cv.get("experience") or []
    if experience:
        entries = []
        for exp in experience:
            role = esc(exp.get("role"))
            company = esc(exp.get("company"))
            dates = esc(exp.get("dates"))
            location = esc(exp.get("location"))
            bullets = exp.get("bullets") or []
            bullets_html = "".join(f"<li>{esc(b)}</li>" for b in bullets)
            title = f"<strong>{role}</strong>, {company}" if role and company else f"<strong>{role or company}</strong>"
            meta = " · ".join(part for part in (dates, location) if part)
            entries.append(f"""
            <div class="entry">
              <div class="entry-header">
                <span class="entry-title">{title}</span>
                <span class="entry-meta">{meta}</span>
              </div>
              {"<ul>" + bullets_html + "</ul>" if bullets_html else ""}
            </div>
            """)
        sections_html.append(f"""
        <section>
          <h2>Experience</h2>
          {''.join(entries)}
        </section>
        """)

    projects = cv.get("projects") or []
    if projects:
        entries = []
        for proj in projects:
            pname = esc(proj.get("name"))
            desc = esc(proj.get("description"))
            entries.append(f"""
            <div class="entry">
              <div class="entry-header">
                <span class="entry-title"><strong>{pname}</strong></span>
              </div>
              {"<p>" + desc + "</p>" if desc else ""}
            </div>
            """)
        sections_html.append(f"""
        <section>
          <h2>Projects</h2>
          {''.join(entries)}
        </section>
        """)

    leadership = cv.get("leadership") or []
    if leadership:
        items = "".join(f"<li>{esc(item)}</li>" for item in leadership)
        sections_html.append(f"""
        <section>
          <h2>Others</h2>
          <ul>{items}</ul>
        </section>
        """)

    skills = cv.get("skills") or {}
    if skills:
        lines = []
        for category, skill_list in skills.items():
            cat = esc(category)
            items = ", ".join(esc(s) for s in (skill_list or []))
            lines.append(f'<div class="skill-line"><span class="skill-category">{cat}:</span> {items}</div>')
        sections_html.append(f"""
        <section>
          <h2>Skills</h2>
          {''.join(lines)}
        </section>
        """)

    body = "".join(sections_html)

    return f"""<!DOCTYPE html>
<html>
<head>
<meta charset="utf-8">
<style>
  body {{
    font-family: Georgia, 'Times New Roman', serif;
    color: #1a1a1a;
    max-width: 800px;
    margin: 0 auto;
    padding: 48px 56px;
    line-height: 1.5;
    background: #ffffff;
  }}
  .header {{ text-align: center; margin-bottom: 28px; }}
  .header h1 {{ font-size: 28px; margin: 0 0 6px; letter-spacing: 0.5px; }}
  .header .contact {{ font-size: 13px; color: #555555; }}
  section {{ margin-bottom: 20px; }}
  h2 {{
    font-size: 13px;
    text-transform: uppercase;
    letter-spacing: 1.5px;
    color: #333333;
    border-bottom: 1px solid #cccccc;
    padding-bottom: 4px;
    margin-bottom: 10px;
  }}
  p {{ font-size: 14px; margin: 4px 0; }}
  .entry {{ margin-bottom: 12px; }}
  .entry-header {{
    display: flex;
    justify-content: space-between;
    align-items: baseline;
    gap: 12px;
    font-size: 14px;
  }}
  .entry-title {{ font-weight: 400; }}
  .entry-meta {{ color: #666666; font-size: 12.5px; white-space: nowrap; }}
  ul {{ margin: 6px 0 0; padding-left: 20px; }}
  li {{ font-size: 13.5px; margin-bottom: 3px; }}
  .skill-line {{ font-size: 13.5px; margin-bottom: 4px; }}
  .skill-category {{ font-weight: 700; }}
</style>
</head>
<body>
  <div class="header">
    <h1>{name}</h1>
    <div class="contact">{contact}</div>
  </div>
  {body}
</body>
</html>"""


def add_section_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.font.small_caps = True
    run.font.size = Pt(13)
    paragraph.paragraph_format.space_before = Pt(10)
    paragraph.paragraph_format.space_after = Pt(2)

    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "000000")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def add_entry(
    doc: Document,
    title: str,
    title_right: str,
    subtitle: str,
    subtitle_right: str,
    bullets: list[str] | None = None,
) -> None:
    title_para = doc.add_paragraph()
    title_para.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
    title_para.paragraph_format.space_after = Pt(2)
    title_run = title_para.add_run(title or "")
    title_run.bold = True
    title_run.font.size = Pt(11)
    title_para.add_run("\t")
    title_right_run = title_para.add_run(title_right or "")
    title_right_run.font.size = Pt(11)

    subtitle_para = doc.add_paragraph()
    subtitle_para.paragraph_format.tab_stops.add_tab_stop(Inches(7.5), WD_TAB_ALIGNMENT.RIGHT)
    subtitle_para.paragraph_format.space_after = Pt(2)
    subtitle_run = subtitle_para.add_run(subtitle or "")
    subtitle_run.italic = True
    subtitle_run.font.size = Pt(10.5)
    subtitle_para.add_run("\t")
    subtitle_right_run = subtitle_para.add_run(subtitle_right or "")
    subtitle_right_run.italic = True
    subtitle_right_run.font.size = Pt(10.5)

    if bullets:
        for bullet_text in bullets:
            bullet_para = doc.add_paragraph(bullet_text, style="List Bullet")
            bullet_para.paragraph_format.space_after = Pt(2)
            for bullet_run in bullet_para.runs:
                bullet_run.font.size = Pt(10.5)


def render_cv_docx(cv: dict) -> Document:
    """Build a CV Document from scratch (no template) matching a fixed visual
    style. Returns the Document; the caller decides where to save it."""
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(11)

    name_para = doc.add_paragraph()
    name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_para.add_run(cv.get("name") or "")
    name_run.bold = True
    name_run.font.small_caps = True
    name_run.font.size = Pt(24)

    contact_para = doc.add_paragraph()
    contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_para.paragraph_format.space_before = Pt(2)
    contact_run = contact_para.add_run(cv.get("contact") or "")
    contact_run.font.size = Pt(10)

    education = cv.get("education") or []
    if education:
        add_section_heading(doc, "Education")
        for entry in education:
            add_entry(
                doc,
                title=entry.get("school", ""),
                title_right=entry.get("location", ""),
                subtitle=entry.get("degree", ""),
                subtitle_right=entry.get("dates", ""),
            )

    experience = cv.get("experience") or []
    if experience:
        add_section_heading(doc, "Experience")
        for entry in experience:
            add_entry(
                doc,
                title=entry.get("role", ""),
                title_right=entry.get("dates", ""),
                subtitle=entry.get("company", ""),
                subtitle_right=entry.get("location", ""),
                bullets=entry.get("bullets") or [],
            )

    projects = cv.get("projects") or []
    if projects:
        add_section_heading(doc, "Projects")
        for project in projects:
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(4)
            project_name_run = para.add_run(f"{project.get('name', '')}: ")
            project_name_run.bold = True
            project_name_run.font.size = Pt(10.5)
            desc_run = para.add_run(project.get("description", ""))
            desc_run.font.size = Pt(10.5)

    leadership = cv.get("leadership") or []
    if leadership:
        add_section_heading(doc, "Others")
        for item in leadership:
            item_para = doc.add_paragraph(item, style="List Bullet")
            for item_run in item_para.runs:
                item_run.font.size = Pt(10.5)

    skills = cv.get("skills") or {}
    if skills:
        add_section_heading(doc, "Skills")
        for category, skill_list in skills.items():
            para = doc.add_paragraph()
            para.paragraph_format.space_after = Pt(2)
            category_run = para.add_run(f"{category}: ")
            category_run.bold = True
            category_run.font.size = Pt(10.5)
            items_run = para.add_run(", ".join(skill_list or []))
            items_run.font.size = Pt(10.5)

    return doc


def make_section_heading(text: str) -> list:
    heading_style = ParagraphStyle(
        name="PdfSectionHeading",
        fontName="Times-Roman",
        fontSize=13,
        leading=16,
        spaceBefore=10,
        spaceAfter=2,
    )
    return [
        PdfParagraph(html.escape(text.upper()), heading_style),
        HRFlowable(width="100%", thickness=0.75, color=black, spaceAfter=6),
    ]


def make_entry(
    title: str,
    title_right: str,
    subtitle: str,
    subtitle_right: str,
    bullets: list[str] | None = None,
) -> list:
    title_style = ParagraphStyle(name="PdfEntryTitle", fontName="Times-Bold", fontSize=11, leading=14)
    title_right_style = ParagraphStyle(name="PdfEntryTitleRight", fontName="Times-Roman", fontSize=11, leading=14, alignment=TA_RIGHT)
    subtitle_style = ParagraphStyle(name="PdfEntrySubtitle", fontName="Times-Italic", fontSize=10.5, leading=13)
    subtitle_right_style = ParagraphStyle(name="PdfEntrySubtitleRight", fontName="Times-Italic", fontSize=10.5, leading=13, alignment=TA_RIGHT)

    table = Table(
        [
            [
                PdfParagraph(html.escape(title or ""), title_style),
                PdfParagraph(html.escape(title_right or ""), title_right_style),
            ],
            [
                PdfParagraph(html.escape(subtitle or ""), subtitle_style),
                PdfParagraph(html.escape(subtitle_right or ""), subtitle_right_style),
            ],
        ],
        colWidths=[5.5 * inch, 1.5 * inch],
    )
    table.setStyle(TableStyle([
        ("LEFTPADDING", (0, 0), (-1, -1), 0),
        ("RIGHTPADDING", (0, 0), (-1, -1), 0),
        ("TOPPADDING", (0, 0), (-1, -1), 0),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
    ]))

    flowables: list = [table]

    if bullets:
        bullet_style = ParagraphStyle(name="PdfBullet", fontName="Times-Roman", fontSize=10.5, leading=13, spaceAfter=2)
        items = [ListItem(PdfParagraph(html.escape(b), bullet_style)) for b in bullets]
        flowables.append(ListFlowable(items, bulletType="bullet", leftIndent=14))

    return flowables


def render_cv_pdf(cv: dict) -> bytes:
    """Build a CV PDF from scratch (no template), mirroring render_cv_docx's
    visual style. Returns raw PDF bytes."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
        leftMargin=0.5 * inch,
        rightMargin=0.5 * inch,
    )

    name_style = ParagraphStyle(name="PdfName", fontName="Times-Bold", fontSize=24, leading=28, alignment=TA_CENTER, spaceAfter=4)
    contact_style = ParagraphStyle(name="PdfContact", fontName="Times-Roman", fontSize=10, leading=13, alignment=TA_CENTER, spaceBefore=2)

    story: list = [
        PdfParagraph(html.escape((cv.get("name") or "").upper()), name_style),
        PdfParagraph(html.escape(cv.get("contact") or ""), contact_style),
    ]

    education = cv.get("education") or []
    if education:
        story.extend(make_section_heading("Education"))
        for entry in education:
            story.extend(make_entry(
                title=entry.get("school", ""),
                title_right=entry.get("location", ""),
                subtitle=entry.get("degree", ""),
                subtitle_right=entry.get("dates", ""),
            ))

    experience = cv.get("experience") or []
    if experience:
        story.extend(make_section_heading("Experience"))
        for entry in experience:
            story.extend(make_entry(
                title=entry.get("role", ""),
                title_right=entry.get("dates", ""),
                subtitle=entry.get("company", ""),
                subtitle_right=entry.get("location", ""),
                bullets=entry.get("bullets") or [],
            ))

    projects = cv.get("projects") or []
    if projects:
        story.extend(make_section_heading("Projects"))
        project_style = ParagraphStyle(name="PdfProject", fontName="Times-Roman", fontSize=10.5, leading=13, spaceAfter=4)
        for project in projects:
            text = f"<b>{html.escape(project.get('name', ''))}</b>: {html.escape(project.get('description', ''))}"
            story.append(PdfParagraph(text, project_style))

    leadership = cv.get("leadership") or []
    if leadership:
        story.extend(make_section_heading("Others"))
        others_style = ParagraphStyle(name="PdfOthersItem", fontName="Times-Roman", fontSize=10.5, leading=13, spaceAfter=2)
        items = [ListItem(PdfParagraph(html.escape(item), others_style)) for item in leadership]
        story.append(ListFlowable(items, bulletType="bullet", leftIndent=14))

    skills = cv.get("skills") or {}
    if skills:
        story.extend(make_section_heading("Skills"))
        skill_style = ParagraphStyle(name="PdfSkill", fontName="Times-Roman", fontSize=10.5, leading=13, spaceAfter=2)
        for category, skill_list in skills.items():
            text = f"<b>{html.escape(category)}</b>: {html.escape(', '.join(skill_list or []))}"
            story.append(PdfParagraph(text, skill_style))

    doc.build(story)
    return buffer.getvalue()


def _insert_paragraph_after(paragraph: Paragraph, text: str, style: str | None = None) -> Paragraph:
    new_element = OxmlElement("w:p")
    paragraph._p.addnext(new_element)
    new_paragraph = Paragraph(new_element, paragraph._parent)
    if style:
        new_paragraph.style = style
    if text:
        new_paragraph.add_run(text)
    return new_paragraph


def _remove_existing_bullets_after(heading: Paragraph) -> None:
    """Strip placeholder bullet paragraphs immediately under a section heading."""
    next_element = heading._p.getnext()
    while next_element is not None and next_element.tag == qn("w:p"):
        next_paragraph = Paragraph(next_element, heading._parent)
        style_name = (next_paragraph.style.name if next_paragraph.style else "") or ""
        if "list" in style_name.lower() or "bullet" in style_name.lower():
            following = next_element.getnext()
            next_element.getparent().remove(next_element)
            next_element = following
        else:
            break


def _write_bullets(doc: Document, heading: Paragraph | None, bullets: list[str]) -> None:
    cleaned_bullets = [b.strip(" \t\n-*") for b in bullets]
    cleaned_bullets = [b for b in cleaned_bullets if b]

    if heading is None:
        heading = doc.add_paragraph()
        heading.text = "Experience"

    _remove_existing_bullets_after(heading)

    anchor = heading
    for bullet in cleaned_bullets:
        anchor = _insert_paragraph_after(anchor, bullet, style="List Bullet")
        anchor.paragraph_format.left_indent = Inches(0.2)


def generate_cv_docx(template_path: str, customized_bullets: list[str], output_path: str) -> str:
    output = Path(output_path)
    output.parent.mkdir(parents=True, exist_ok=True)

    doc = load_template_document(template_path) if template_path.lower().endswith(".docx") else Document()
    is_blank_template = len(doc.paragraphs) == 0

    if is_blank_template:
        doc.add_heading("Customized CV", level=1)
        heading = doc.add_paragraph("Experience")
        _write_bullets(doc, heading, customized_bullets)
        doc.save(str(output))
        return str(output)

    target = None
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip().lower()
        if "work experience" in text or "professional experience" in text or "experience" in text:
            target = paragraph
            break

    _write_bullets(doc, target, customized_bullets)

    doc.save(str(output))
    return str(output)


def generate_cover_letter(
    user_profile_json: dict,
    parsed_job: dict,
    template_path: str,
    output_path: str | None = None,
) -> str:
    output = Path(output_path) if output_path else Path(f"/tmp/cover_letter_{uuid.uuid4().hex[:8]}.docx")
    output.parent.mkdir(parents=True, exist_ok=True)

    template_doc = load_template_document(template_path) if template_path.lower().endswith(".docx") else Document()
    template_text = load_template_text(template_path)

    prompt = build_cover_letter_customization_prompt(user_profile_json, parsed_job, template_text)
    try:
        raw_text = call_claude(prompt, max_tokens=1800)
        content = _clean_text_payload(raw_text)
    except ClaudeAPIError as exc:
        logger.warning("generate_cover_letter falling back to deterministic letter: %s", exc)
        content = _build_fallback_cover_letter(user_profile_json, parsed_job)

    paragraphs = list(template_doc.paragraphs)
    if not paragraphs:
        doc = Document()
        doc.add_paragraph(content)
        doc.save(str(output))
        return str(output)

    body_start = None
    body_end = None
    for idx, paragraph in enumerate(paragraphs):
        text = paragraph.text.strip().lower()
        if body_start is None and any(marker in text for marker in ["dear", "hello", "to whom"]):
            body_start = idx + 1
            continue
        if body_end is None and any(marker in text for marker in ["sincerely", "best", "regards", "kind regards", "thanks", "thank you", "signature"]):
            body_end = idx
            break

    if body_start is None:
        body_start = 0
    if body_end is None:
        body_end = len(paragraphs)

    doc = Document()
    for idx in range(0, body_start):
        _copy_paragraph(paragraphs[idx], doc)

    for paragraph_text in [part.strip() for part in content.split("\n\n") if part.strip()]:
        doc.add_paragraph(paragraph_text)

    for idx in range(body_end, len(paragraphs)):
        _copy_paragraph(paragraphs[idx], doc)

    doc.save(str(output))
    return str(output)


def generate_cold_email(user_profile_json: dict, parsed_job: dict) -> str:
    """Never raises: any Claude failure falls back to a deterministic, personalized email."""
    prompt = build_cold_email_prompt(user_profile_json, parsed_job)
    try:
        raw_text = call_claude(prompt, max_tokens=800)
        text = _clean_text_payload(raw_text)
        if text.startswith("Subject:"):
            text = "\n".join(text.splitlines()[1:]) if len(text.splitlines()) > 1 else ""
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        if not lines:
            return _build_fallback_cold_email(user_profile_json, parsed_job)
        return "\n".join(lines)
    except ClaudeAPIError as exc:
        logger.warning("generate_cold_email falling back to deterministic email: %s", exc)
        return _build_fallback_cold_email(user_profile_json, parsed_job)
